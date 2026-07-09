#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Render a Thai meeting-summary Markdown file to .docx.

Usage:
    python3 gen_docx.py <input.md> [output.docx]

If output is omitted, writes alongside the input with a .docx extension.
Font: TH SarabunPSK, body 16pt, headers scaled, Arabic numerals preserved.
Requires: python-docx
"""

import re
import sys
import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "TH SarabunPSK"
GLYPH_FONT = "DejaVu Sans"  # has ● ✓ ✗ (emoji fonts tofu in LibreOffice)

# Icon emoji -> (glyph, RGB, bold). Rendered as a colored glyph so status /
# severity markers show in real color in both Word and LibreOffice, instead of
# tofu'd emoji. Keeps the .md readable (emoji) while the .docx gets colored dots.
ICONS = {
    '🔴': ('●', (192, 57, 43), False),   # red    — high / critical
    '🟠': ('●', (230, 126, 34), False),  # orange — medium
    '🟡': ('●', (194, 149, 0), False),   # yellow — low
    '🟢': ('●', (39, 174, 96), False),   # green  — ok
    '🔵': ('●', (46, 134, 193), False),  # blue   — info
    '✅': ('✓', (39, 174, 96), True),    # green check
    '✔': ('✓', (39, 174, 96), True),
    '✓': ('✓', (39, 174, 96), True),
    '❌': ('✗', (192, 57, 43), True),    # red cross
    '✗': ('✗', (192, 57, 43), True),
}
_ICON_RE = re.compile('(' + '|'.join(re.escape(k) for k in ICONS) + ')')


def set_run_font(run, size=16, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:cs'), font)


def _emit(p, text, size, bold, color):
    """Emit runs for text, turning icon emoji into colored glyphs."""
    for seg in _ICON_RE.split(text):
        if not seg:
            continue
        icon = ICONS.get(seg)
        if icon:
            glyph, gcolor, gbold = icon
            set_run_font(p.add_run(glyph), size, gbold, gcolor, font=GLYPH_FONT)
        else:
            set_run_font(p.add_run(seg), size, bold, color)


def add_runs(p, text, size=16, bold=False, color=None):
    """Add text to paragraph, honoring inline **bold** and colored icon glyphs."""
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _emit(p, part[2:-2], size, True, color)
        else:
            _emit(p, part, size, bold, color)


def add_para(doc, text, size=16, bold=False, align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    add_runs(p, text, size, bold, color)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ''
        set_run_font(cell.paragraphs[0].add_run(h), 15, True, (255, 255, 255))
        shade_cell(cell, '2E5C8A')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i >= len(cells):
                break
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, val, 14)
    return t


def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def parse_table_block(lines, idx):
    header = split_row(lines[idx])
    idx += 2  # skip the |---|---| separator
    rows = []
    while idx < len(lines) and lines[idx].lstrip().startswith('|'):
        rows.append(split_row(lines[idx]))
        idx += 1
    return header, rows, idx


def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '1'), ('w:color', 'BBBBBB')):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def main():
    if len(sys.argv) < 2:
        sys.exit("Usage: gen_docx.py <input.md> [output.docx]")
    md_path = os.path.abspath(sys.argv[1])
    if not os.path.isfile(md_path):
        sys.exit(f"Input not found: {md_path}")

    if len(sys.argv) > 2:
        # Defense-in-depth: keep output basename inside the input's directory,
        # ignoring any path components in the supplied name (CWE-22).
        out_path = os.path.join(os.path.dirname(md_path),
                                os.path.basename(sys.argv[2]))
    else:
        out_path = os.path.splitext(md_path)[0] + ".docx"

    with open(md_path, encoding='utf-8', errors='replace') as f:
        lines = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(16)
    normal.element.rPr.rFonts.set(qn('w:cs'), FONT)

    BLUE_D, BLUE = (31, 61, 92), (46, 92, 138)
    GREY = (90, 90, 90)

    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s == '':
            i += 1
            continue
        if s == '---':
            hr(doc); i += 1; continue
        if s.startswith('|'):
            header, rows, i = parse_table_block(lines, i)
            add_table(doc, header, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if s.startswith('#### '):
            add_para(doc, s[5:], 16, True, color=BLUE, space_after=4)
        elif s.startswith('### '):
            add_para(doc, s[4:], 17, True, color=BLUE, space_after=4)
        elif s.startswith('## '):
            add_para(doc, s[3:], 19, True, color=BLUE_D, space_after=6)
        elif s.startswith('# '):
            add_para(doc, s[2:], 24, True, WD_ALIGN_PARAGRAPH.CENTER, 8, BLUE_D)
        elif s.startswith('> '):
            p = add_para(doc, s[2:], 14, color=GREY, space_after=6)
            p.paragraph_format.left_indent = Pt(14)
        elif s.startswith('- '):
            p = add_para(doc, s[2:], 16, space_after=3)
            p.paragraph_format.left_indent = Pt(20)
            p.style = doc.styles['List Bullet']
            for r in p.runs:
                set_run_font(r, 16, r.bold)
        elif len(s) > 2 and s[0].isdigit() and s[1] == '.':
            p = add_para(doc, s[s.index('.') + 1:].strip(), 16, space_after=3)
            p.paragraph_format.left_indent = Pt(20)
        elif s.startswith('*') and s.endswith('*') and not s.startswith('**'):
            add_para(doc, s.strip('*'), 14, color=GREY, space_after=6)
        else:
            add_para(doc, s, 16, space_after=4)
        i += 1

    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == '__main__':
    main()
